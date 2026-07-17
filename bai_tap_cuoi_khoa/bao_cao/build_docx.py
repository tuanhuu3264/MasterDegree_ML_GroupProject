# -*- coding: utf-8 -*-
"""Dung bao cao Word (.docx) tu noi dung bao cao hoc thuat (khop PDF).
Chay: python bao_cao/build_docx.py  (tu thu muc goc repo)
Yeu cau: python-docx, cac anh trong bao_cao/assets/"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = 'bao_cao/assets/'
NAVY = RGBColor(0x14, 0x24, 0x3d)
BLUE = RGBColor(0x1f, 0x3a, 0x5f)
GRAY = RGBColor(0x55, 0x55, 0x55)
BODY = 'Cambria'
SANS = 'Calibri'

doc = Document()

# ---- Page A4 + margins ----
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)

# ---- Default style ----
st = doc.styles['Normal']
st.font.name = BODY; st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.25

# ================= HELPERS =================
def _set_font(run, name=BODY, size=11, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color is not None: run.font.color.rgb = color
    # dam bao font cho tieng Viet
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), name)

def para(text='', size=11, bold=False, italic=False, color=None, align=None,
         space_before=None, space_after=None, justify=True, font=BODY):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    elif justify: p.alignment = AL.JUSTIFY
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); _set_font(r, font, size, bold, italic, color)
    return p

def runs_para(runs, align=AL.JUSTIFY, space_after=None):
    """runs = list of (text, dict)"""
    p = doc.add_paragraph(); p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    for text, kw in runs:
        r = p.add_run(text); _set_font(r, **kw)
    return p

def heading(text, level=2):
    if level == 2:
        p = para(space_before=14, space_after=6, justify=False)
        r = p.add_run(text); _set_font(r, BODY, 14, True, False, NAVY)
        # duong ke duoi
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '14243D')
        pbdr.append(bottom); pPr.append(pbdr)
    elif level == 3:
        p = para(space_before=10, space_after=4, justify=False)
        r = p.add_run(text); _set_font(r, BODY, 12, True, False, BLUE)
    elif level == 4:
        p = para(space_before=8, space_after=3, justify=False)
        r = p.add_run(text); _set_font(r, BODY, 11, True, True, NAVY)
    return p

def caption(text_bold, text_rest):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text_bold); _set_font(r, SANS, 9.5, True, True, NAVY)
    r2 = p.add_run(text_rest); _set_font(r2, SANS, 9.5, False, True, GRAY)
    return p

def note(label, text):
    p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    # vien trai
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '10'); left.set(qn('w:color'), '1F3A5F')
    pbdr.append(left); pPr.append(pbdr)
    r = p.add_run(label + ' '); _set_font(r, BODY, 10.5, True, False, NAVY)
    r2 = p.add_run(text); _set_font(r2, BODY, 10.5, False, False, None)
    return p

def set_cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)

def _cell_text(cell, text, bold=False, size=9.5, color=None, align=AL.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text); _set_font(r, SANS, size, bold, False, color)

def table(cap_bold, cap_rest, headers, rows, widths=None):
    caption(cap_bold, cap_rest)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; _cell_text(c, h, bold=True, size=9.5, color=NAVY)
        set_cell_bg(c, 'E8ECF2')
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            _cell_text(t.rows[i].cells[j], val, size=9.3)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, cap_bold, cap_rest, width_in):
    if not os.path.exists(path):
        para(f'[thiếu ảnh: {path}]', italic=True); return
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width_in))
    cp = doc.add_paragraph(); cp.alignment = AL.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(cap_bold); _set_font(r, SANS, 9.3, True, False, NAVY)
    r2 = cp.add_run(cap_rest); _set_font(r2, SANS, 9.3, False, False, GRAY)

def pagebreak():
    doc.add_page_break()

def blank(n=1, size=11):
    for _ in range(n):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(' '); _set_font(r, BODY, size)

# ================= TRANG BIA =================
def cover_line(text, size, bold=False, color=None, sb=0, sa=0, italic=False):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); _set_font(r, BODY, size, bold, italic, color)
    return p

cover_line('MINISTRY OF EDUCATION AND TRAINING', 15, True, NAVY, sb=6)
cover_line('FPT UNIVERSITY', 15, True, NAVY, sa=6)
blank(6)
cover_line('Predictive Maintenance of CNC Machines under Distribution Shift', 17, False, NAVY, sa=4)
cover_line('Dự đoán hỏng hóc thiết bị CNC trong điều kiện dịch chuyển phân phối dữ liệu', 11, False, GRAY, italic=True)
blank(3)
cover_line('by', 13, False, NAVY)
blank(1)
for ms, nm in [('25MS23315','Trịnh Hữu Tuấn'),('25MS23323','Hoàng Phi Hải'),
               ('25MS23327','Đỗ Hoàng Tỷ Phú'),('25MS23328','Lê Lâm Vĩnh')]:
    cover_line(f'{ms}   ·   {nm}', 12.5, False, None, sa=2)
blank(1)
cover_line('Supervisor: Dr. Cao Tiến Dũng', 12, False, NAVY)
blank(3)
cover_line('A group project submitted in conformity with the requirements', 11.5)
cover_line('for the fulfillment of the course "Machine Learning"', 11.5)
blank(2)
cover_line('© Copyright by Group of 4 students   ·   2026', 11.5)
pagebreak()

# ================= TOM TAT =================
heading('Tóm tắt', 2)
runs_para([
    ('Tóm tắt. ', dict(name=BODY, size=10.5, bold=True, italic=True, color=NAVY)),
    ('Báo cáo trình bày một phương pháp dự đoán hỏng hóc thiết bị CNC ở ca vận hành kế tiếp, phát biểu dưới dạng bài '
     'toán phân loại nhị phân mất cân bằng (tỉ lệ hỏng khoảng 8%). Thách thức trọng tâm không nằm ở việc lựa chọn '
     'thuật toán mà ở hiện tượng dịch chuyển phân phối (distribution shift): mô hình được huấn luyện trên Dây chuyền A '
     'song phải vận hành trên Dây chuyền B có điều kiện nhiệt và tốc độ khác biệt, trong khi nhãn của tập kiểm thử '
     'không được sử dụng cho quá trình lựa chọn mô hình. Nghiên cứu chẩn đoán hiện tượng này là dịch chuyển hiệp biến '
     '(covariate shift) [4] và đề xuất bốn đặc trưng "khoảng cách tới biên" neo trên các ngưỡng của cơ chế sinh lỗi '
     'được khôi phục từ chính Dây chuyền A (HDF, PWF, OSF, TWF) — theo tinh thần bộ dữ liệu AI4I 2020 [1]. Vì các '
     'ngưỡng này là thuộc tính cơ chế bất biến giữa hai dây chuyền, ranh giới quyết định học trên Dây chuyền A chuyển '
     'giao trực tiếp sang Dây chuyền B (PSI ≈ 0; Drift-AUC của đặc trưng biên = 0,53). Các kỹ thuật hiệu chỉnh trọng '
     'số theo tỉ lệ mật độ [2], hiệu chỉnh ngưỡng và học kết hợp được áp dụng như các lớp tinh chỉnh, toàn bộ được lựa '
     'chọn thông qua kiểm định có trọng số quan trọng (Importance-Weighted Validation) [3] nhằm bảo đảm không rò rỉ '
     'nhãn tập kiểm thử. Mô hình cuối cùng (học kết hợp bỏ phiếu giữa Random Forest, ExtraTrees và XGBoost) đạt '
     'F1 = 0,783 trên Dây chuyền B, cải thiện 3,4 lần so với mô hình cơ sở. Phân tích cho thấy con số này chạm trần '
     'hiệu năng của dữ liệu: khoảng 25% ca hỏng là nhiễu ngẫu nhiên không thể học được.',
     dict(name=BODY, size=10.5))])
runs_para([
    ('Từ khóa: ', dict(name=BODY, size=10, bold=True, italic=True, color=NAVY)),
    ('bảo trì dự đoán; dịch chuyển hiệp biến; đặc trưng bất biến; hiệu chỉnh trọng số theo tỉ lệ mật độ; học kết hợp; '
     'dữ liệu mất cân bằng.', dict(name=BODY, size=10))])
pagebreak()

# ================= MUC LUC =================
heading('Mục lục', 2)
toc = [
    ('1. Giới thiệu', []),
    ('2. Dữ liệu và phát biểu bài toán', ['2.1 Mô tả dữ liệu','2.2 Chẩn đoán loại dịch chuyển phân phối','2.3 Lựa chọn độ đo đánh giá']),
    ('3. Phân tích khám phá dữ liệu', []),
    ('4. Tổng hợp các quyết định thiết kế', []),
    ('5. Phương pháp đề xuất và luận cứ', ['5.1 Độ đo đánh giá','5.2 Chiến lược chia dữ liệu',
        '5.3 Thiết kế đặc trưng (khôi phục ngưỡng luật sinh nhãn từ Dây chuyền A)','5.4 Cơ chế ngưỡng và tính bất biến hạng',
        '5.5 Lựa chọn họ mô hình','5.6 Hiệu chỉnh dịch chuyển bằng trọng số tỉ lệ mật độ',
        '5.7 Kiểm định có trọng số quan trọng (IWV)','5.8 Học kết hợp','5.9 Ngưỡng quyết định']),
    ('6. Thực nghiệm và kết quả', ['6.1 Đánh giá độ tin cậy bằng bootstrap']),
    ('7. Thảo luận: hàm ý vận hành và bảo trì', []),
    ('8. Hạn chế của phương pháp', []),
    ('9. Hướng phát triển', []),
    ('10. Kết luận', []),
]
for main, subs in toc:
    p = para(space_after=3, justify=False)
    r = p.add_run(main); _set_font(r, BODY, 11.5, True, False, NAVY)
    for sub in subs:
        sp = para(space_after=1, justify=False)
        sp.paragraph_format.left_indent = Cm(0.8)
        r = sp.add_run(sub); _set_font(r, BODY, 10.2, False, False, GRAY)
p = para(space_before=8, justify=False)
r = p.add_run('Phụ lục.  Kiểm soát rò rỉ dữ liệu'); _set_font(r, BODY, 11, True, False, NAVY)
p = para(space_after=1, justify=False)
r = p.add_run('Tài liệu tham khảo'); _set_font(r, BODY, 11, True, False, NAVY)
pagebreak()

# ================= 1. GIOI THIEU =================
heading('1. Giới thiệu', 2)
para('Bảo trì dự đoán cho phép can thiệp vào thiết bị trước khi sự cố xảy ra, qua đó giảm thiểu thời gian dừng dây '
     'chuyền ngoài kế hoạch. Bài toán được phát biểu dưới dạng phân loại nhị phân với nhãn hong_hoc ∈ {0,1}, trong đó '
     'nhãn dương biểu thị thiết bị sẽ hỏng ở ca kế tiếp. Trong bối cảnh vận hành, sai lầm bỏ sót (dự báo âm cho một '
     'thiết bị thực tế hỏng) gây tổn thất lớn hơn báo động giả, do đó phương pháp ưu tiên độ nhạy (Recall) nhưng vẫn '
     'cân bằng với độ chính xác (Precision) thông qua độ đo F1.')
para('Điểm khác biệt của bài toán so với một bài phân loại thông thường là yêu cầu chuyển giao mô hình giữa hai phân '
     'phối dữ liệu: tập huấn luyện thu thập từ Dây chuyền A (nhà máy cũ) trong khi tập kiểm thử đến từ Dây chuyền B '
     '(nhà máy mới, điều kiện nhiệt và tốc độ cao hơn). Ràng buộc quan trọng là nhãn của Dây chuyền B không được sử '
     'dụng trong quá trình lựa chọn mô hình hay hiệu chỉnh siêu tham số. Toàn bộ báo cáo được tổ chức theo hướng luận '
     'giải: mỗi quyết định kỹ thuật đều nêu rõ căn cứ lựa chọn, các phương án thay thế đã cân nhắc và bị loại bỏ, '
     'cùng rủi ro trong trường hợp giả định nền tảng không được thỏa mãn.')

# ================= 2. DU LIEU =================
heading('2. Dữ liệu và phát biểu bài toán', 2)
heading('2.1 Mô tả dữ liệu', 3)
table('Bảng 1. ', 'Đặc điểm bộ dữ liệu.', ['Thuộc tính', 'Chi tiết'], [
    ['Kích thước', 'Train A: 14.000 × 8; Test B: 6.000 × 8; không có giá trị thiếu, không có bản ghi trùng lặp.'],
    ['Biến số (5)', 'nhiệt độ môi trường, nhiệt độ quy trình, tốc độ quay, mômen xoắn, độ mòn dao.'],
    ['Biến phân loại (2)', 'loai_san_pham (L < M < H, có thứ tự); ca_lam_viec (danh nghĩa).'],
    ['Nhãn', 'hong_hoc; tỉ lệ hỏng: 7,36% (A) và 7,95% (B).'],
], widths=[4.2, 12.5])

heading('2.2 Chẩn đoán loại dịch chuyển phân phối', 3)
para('Việc xác định đúng loại dịch chuyển là điều kiện tiên quyết, bởi mỗi loại đòi hỏi một chiến lược xử lý khác '
     'nhau [4]. Bảng 2 trình bày ba loại dịch chuyển và bằng chứng tương ứng trong dữ liệu.')
table('Bảng 2. ', 'Chẩn đoán loại dịch chuyển phân phối.',
    ['Loại dịch chuyển', 'Đặc điểm', 'Bằng chứng trong dữ liệu'], [
    ['Dịch chuyển hiệp biến — P(x) thay đổi, P(y|x) ổn định', 'Phân phối đầu vào thay đổi, cơ chế sinh nhãn giữ nguyên',
     'Có: biến thô dịch mạnh nhưng tỉ lệ nhãn ổn định (7,36% → 7,95%); P(hỏng | vượt ngưỡng) gần như không đổi (mục 5.4)'],
    ['Dịch chuyển nhãn — P(y) thay đổi', 'Tỉ lệ lớp thay đổi đáng kể', 'Không: tiên nghiệm nhãn gần như không đổi'],
    ['Dịch chuyển khái niệm — P(y|x) thay đổi', 'Cùng đầu vào nhưng quan hệ với nhãn thay đổi',
     'Không quan sát thấy; tuy nhiên chưa thể loại trừ tuyệt đối do thiếu nhãn B (xem mục 8)'],
])
para('Kết luận về dịch chuyển hiệp biến cho phép áp dụng bộ công cụ phù hợp: đặc trưng bất biến, hiệu chỉnh trọng số '
     'theo tỉ lệ mật độ, và kiểm định có trọng số quan trọng. Nếu chẩn đoán sai thành dịch chuyển khái niệm, các kỹ '
     'thuật này sẽ không còn hiệu lực.')

heading('2.3 Lựa chọn độ đo đánh giá', 3)
note('Về việc loại bỏ độ chính xác (accuracy).', 'Với tỉ lệ hỏng khoảng 8%, một mô hình dự báo toàn bộ là "không '
     'hỏng" đạt độ chính xác 92% nhưng độ nhạy bằng 0. Độ chính xác thưởng cho việc dự đoán đúng lớp đa số — đối tượng '
     'không phải là mối quan tâm của bài toán — nên bị loại khỏi hệ độ đo.')
para('Về ưu tiên AUC-PR so với AUC-ROC: AUC-ROC được tính trên nền toàn bộ lớp âm (5.523 mẫu), do đó dễ cho giá trị '
     'cao một cách thiếu ý nghĩa khi lớp âm áp đảo. Ngược lại, AUC-PR chỉ đánh giá chất lượng của các cảnh báo dương, '
     'phù hợp với mục tiêu vận hành và được khuyến nghị cho dữ liệu mất cân bằng mạnh [10]. Báo cáo trình bày đồng '
     'thời ba độ đo: F1 là độ đo so sánh chính, AUC-PR là tham chiếu về độ nhạy với lớp hiếm, và AUC-ROC dùng để so '
     'sánh khả năng xếp hạng.')

# ================= 3. EDA =================
heading('3. Phân tích khám phá dữ liệu', 2)
para('Phân tích khám phá được thực hiện nhằm định hướng quyết định thiết kế. Bốn phát hiện chính và hệ quả kỹ thuật '
     'tương ứng được tổng hợp trong Bảng 3.')
table('Bảng 3. ', 'Các phát hiện khám phá và quyết định kéo theo.', ['#', 'Phát hiện (bằng chứng)', 'Quyết định kéo theo'], [
    ['1', 'Mất cân bằng lớp khoảng 8%, ổn định giữa A và B', 'Loại độ chính xác, dùng F1/AUC-PR (§5.1); là dịch chuyển hiệp biến nên phù hợp hiệu chỉnh trọng số (§5.6)'],
    ['2', 'Nhiệt độ và tốc độ trên B dịch mạnh, độ phân tán rộng hơn khoảng 30% (Hình 1)', 'Không sử dụng biến thô làm ranh giới quyết định; cần đặc trưng bất biến (§5.3)'],
    ['3', 'Độ mòn dao là tín hiệu mạnh nhất (hệ số tương quan +0,20) và ổn định (PSI = 0,001)', 'Chỉ báo dự báo then chốt, đồng thời can thiệp được; cơ sở cho khuyến nghị bảo trì (§7)'],
    ['4', 'Tín hiệu hỏng mang tính phi tuyến, tập trung ở các đuôi phân phối (mòn dao cao, tốc độ thấp, mômen ở cả hai biên)', 'Cần mô hình cây; mô hình tuyến tính thuần không đủ (§5.5)'],
], widths=[0.8, 8.0, 7.9])
figure(ASSETS+'02_shift.png', 'Hình 1. ', 'Phân phối chồng lấn giữa Dây chuyền A và B: nhiệt độ và tốc độ dịch chuyển rõ rệt, trong khi độ mòn dao ổn định.', 6.3)
note('Phát hiện bổ sung — rủi ro ngoại suy.', 'Khoảng 2,8% số mẫu trên B nằm ngoài miền giá trị quan sát được của A '
     '(chẳng hạn tốc độ cực đại của B là 2.414 so với 2.153 của A). Do mô hình cây không có khả năng ngoại suy, phát '
     'hiện này trực tiếp dẫn tới quyết định giữ mô hình hồi quy logistic trong tổ hợp học kết hợp như một cơ chế phòng '
     'vệ (§5.5, §5.8).')

# ================= 4. QUYET DINH =================
pagebreak()
heading('4. Tổng hợp các quyết định thiết kế', 2)
para('Bảng 4 tổng hợp toàn bộ quyết định kỹ thuật của nghiên cứu; mỗi hàng nêu căn cứ lựa chọn, phương án thay thế đã '
     'bị loại bỏ và rủi ro nếu giả định nền tảng không được thỏa mãn. Luận cứ chi tiết được trình bày trong Mục 5.')
table('Bảng 4. ', 'Tổng hợp quyết định thiết kế và luận cứ.',
    ['Quyết định', 'Căn cứ lựa chọn', 'Phương án thay thế bị loại bỏ', 'Rủi ro nếu giả định sai'], [
    ['Độ đo: F1 và AUC-PR', 'Lớp mất cân bằng khiến độ chính xác vô nghĩa; bỏ sót đắt hơn báo động giả', 'Độ chính xác (thưởng lớp đa số); chỉ tối đa Recall (dự báo hỏng toàn bộ)', 'Thấp — độ đo là ràng buộc của đề bài'],
    ['Chia dữ liệu theo dây chuyền A→B', 'Phản ánh đúng bản chất dịch chuyển và điều kiện triển khai thực tế', 'Chia ngẫu nhiên/phân tầng — trộn lẫn A và B, làm rò rỉ thông tin B', 'Thấp — do đề bài quy định'],
    ['Cân bằng lớp bằng trọng số, không sinh mẫu giả', 'Chia theo dây chuyền làm mất phân tầng; trọng số bù không tạo dữ liệu giả', 'SMOTE/lấy mẫu tăng — nội suy mẫu giả, rủi ro dưới dịch chuyển', 'Trung bình — nếu 8% là quá ít, trọng số có thể chưa đủ'],
    ['Đặc trưng khoảng cách tới biên vật lý', 'Ngưỡng vật lý bất biến giúp ranh giới chuyển giao A→B', 'Tương tác thô/đa thức (dịch theo phân phối A); để cây tự học ngưỡng', 'Cao — nếu ngưỡng vật lý sai lệch, đặc trưng mất giá trị'],
    ['Mã hóa: thứ tự cho loại SP, one-hot cho ca', 'Loại sản phẩm có thứ tự thực; ca làm việc là danh nghĩa', 'One-hot cho loại SP (bỏ thông tin thứ tự); mã thứ tự cho ca (áp đặt thứ tự giả)', 'Thấp'],
    ['Họ mô hình: cây và hồi quy logistic', 'Tín hiệu phi tuyến; cây vượt trội trên dữ liệu bảng; logistic ngoại suy được', 'Chỉ tuyến tính (kém rõ rệt); học sâu (dư thừa, dễ quá khớp)', 'Thấp'],
    ['Tối ưu siêu tham số: RandomizedSearchCV theo AUC-PR', 'Không gian lớn nên tìm ngẫu nhiên hiệu quả hơn tìm lưới', 'GridSearch (vét cạn, kém hiệu quả); tối ưu theo độ chính xác (sai độ đo)', 'Thấp'],
    ['Hiệu chỉnh dịch chuyển bằng trọng số tỉ lệ mật độ', 'Điều chỉnh phân phối huấn luyện về phía B mà không cần nhãn B', 'Bỏ qua dịch chuyển (ranh giới lệch); huấn luyện lại trên B (không có nhãn B)', 'Trung bình — sai nếu không phải dịch chuyển hiệp biến thuần'],
    ['Lựa chọn mô hình/ngưỡng bằng IWV', 'Bài toán chấm trên B nhưng cấm sử dụng nhãn B', 'Lựa chọn theo điểm số trên B (rò rỉ nhãn kiểm thử)', 'Cao — cỡ mẫu hiệu dụng thấp (24,2%) khiến IWV nhiễu'],
    ['Học kết hợp: chọn bỏ phiếu thay vì xếp chồng', 'Bốn mô hình cơ sở ít tương quan; bỏ phiếu đơn giản, ít quá khớp', 'Xếp chồng — dồn trọng số vào XGBoost, kém ổn định hơn dưới dịch chuyển', 'Thấp — hai phương án gần tương đương (F1-IWV bằng nhau)'],
], widths=[2.9, 4.6, 5.0, 4.2])

print('... da xong nua dau, tiep tuc phan 5-references')

# ================= 5. PHUONG PHAP =================
pagebreak()
heading('5. Phương pháp đề xuất và luận cứ', 2)
heading('5.1 Độ đo đánh giá', 3)
para('Như đã lập luận ở Mục 2.3, độ chính xác bị loại và AUC-PR được ưu tiên so với AUC-ROC. Độ đo F1 được chọn làm '
     'tiêu chí so sánh đơn trị vì là trung bình điều hòa của Precision và Recall, phạt nặng khi một trong hai thành '
     'phần thấp — phù hợp yêu cầu vừa hạn chế bỏ sót vừa hạn chế báo động giả. Do lớp hiếm chỉ gồm 477 mẫu trên B '
     'khiến ước lượng F1 có phương sai lớn, khoảng tin cậy 95% được ước lượng bằng phương pháp bootstrap (Mục 6.1).')

heading('5.2 Chiến lược chia dữ liệu', 3)
note('Chia theo dây chuyền thay vì ngẫu nhiên.', 'Mục tiêu thực tế là mô hình huấn luyện trên nhà máy cũ vận hành '
     'được trên nhà máy mới. Nếu trộn A và B rồi chia ngẫu nhiên, tập kiểm định sẽ chứa mẫu của B, khiến mô hình tiếp '
     'xúc trước với phân phối B, làm điểm kiểm định cao giả tạo và triệt tiêu khả năng đo lường dịch chuyển. Chia theo '
     'dây chuyền bảo toàn đúng ranh giới triển khai thực tế.')
para('Về xử lý mất cân bằng, phương pháp không sử dụng SMOTE hay lấy mẫu tăng. Các kỹ thuật này nội suy mẫu hỏng giả '
     'từ các mẫu hỏng lân cận [9]; dưới điều kiện dịch chuyển, vùng lớp hiếm của A không đại diện cho B, khiến mẫu giả '
     'làm mô hình tự tin sai lệch. Thay vào đó, phương pháp sử dụng class_weight=\'balanced\' cho hồi quy logistic và '
     'Random Forest, cùng scale_pos_weight ≈ 12,6 cho XGBoost. Toàn bộ bước chuẩn hóa và mã hóa được đóng gói trong '
     'Pipeline của thư viện scikit-learn [11] và huấn luyện lại trên tập huấn luyện ở mỗi vòng kiểm định chéo nhằm '
     'ngăn rò rỉ dữ liệu.')

heading('5.3 Thiết kế đặc trưng', 3)
para('Phân tích khám phá cho thấy tín hiệu hỏng mang tính phi tuyến. Nhiều phương án tạo phi tuyến đã được cân nhắc '
     'và đánh giá như trong Bảng 5.')
table('Bảng 5. ', 'So sánh các phương án tạo đặc trưng phi tuyến.', ['Phương án', 'Đánh giá'], [
    ['Tương tác thô (mòn × mômen), đa thức bậc hai', 'Bị loại — được dựng từ biến thô nên dịch chuyển theo phân phối A, không hỗ trợ chuyển giao'],
    ['Để mô hình cây tự học ngưỡng phân tách', 'Đúng một phần, nhưng cây học ngưỡng tối ưu trên A; khi B dịch chuyển, ngưỡng lệch vị trí'],
    ['Khoảng cách tới biên vật lý (lựa chọn)', 'Được chọn — neo trên hằng số vật lý không đổi giữa A và B, do đó ranh giới bất biến'],
], widths=[5.5, 11.2])
para('Bốn đặc trưng đề xuất, mỗi đặc trưng mã hóa một cơ chế sinh lỗi dưới dạng "khoảng cách tới biên nguy hiểm", '
     'được mô tả trong Bảng 6.')
table('Bảng 6. ', 'Bốn đặc trưng khoảng cách tới biên và luận cứ dạng hàm.', ['Đặc trưng', 'Công thức', 'Cơ chế và luận cứ dạng hàm'], [
    ['nguy_tan_nhiet', 'max(8,6 − ΔT, 0) × max(1380 − tốc độ, 0)', 'Cơ chế HDF (tản nhiệt). Tích hai hàm bản lề tương đương phép hội (AND): chỉ nguy hiểm khi đồng thời chênh lệch nhiệt nhỏ và tốc độ thấp.'],
    ['lech_cong_suat', 'max(2600 − P, 0) + max(P − 11500, 0)', 'Cơ chế PWF (công suất). Tổng hai hàm bản lề tương đương phép tuyển (OR): nguy hiểm khi công suất ra ngoài dải an toàn [2600, 11500] W.'],
    ['bien_overstrain', 'mòn × mômen − ngưỡng(loại)  {L:12800; M:13900; H:14500}', 'Cơ chế OSF (quá tải). Biên có dấu theo loại sản phẩm; không dùng hàm bản lề vì thông tin "còn xa ngưỡng" cũng hữu ích.'],
    ['mon_twf', 'max(mòn − 244, 0)', 'Cơ chế TWF (mòn dao). Hàm bản lề vượt ngưỡng thay dao; là luật mà cấu hình gốc bỏ sót — hệ số tương quan với nhãn đạt 0,43, mạnh nhất.'],
], widths=[3.0, 5.4, 8.3])
heading('Khôi phục ngưỡng của luật sinh nhãn từ Dây chuyền A', 4)
para('Một câu hỏi then chốt là các hằng số ngưỡng ở Bảng 6 được xác định như thế nào. Bộ dữ liệu sinh theo cơ chế '
     'kiểu AI4I [1] tuân theo các luật có ngưỡng xác định; do đó, với mỗi cơ chế, nghiên cứu quét thử ngưỡng chỉ trên '
     'Dây chuyền A và xác định điểm mà xác suất hỏng có điều kiện tăng vọt — đó chính là biên mà luật sử dụng để sinh '
     'nhãn. Toàn bộ quá trình không sử dụng nhãn của B; các biên sau đó được kiểm chứng là bất biến khi chuyển sang B.')
table('Bảng 6b. ', 'Căn cứ khôi phục ngưỡng từ Dây chuyền A (không dùng nhãn B).', ['Cơ chế', 'Ngưỡng chốt', 'Bằng chứng trên A'], [
    ['HDF', 'ΔT < 8,6 K và tốc độ < 1380', 'Giữ nguyên hằng số AI4I: P(hỏng) tại đây đạt 0,84; nới lỏng làm độ chính xác rớt còn ~0,45.'],
    ['PWF', 'công suất ∉ [2600; 11500] W', 'Thiết bị không hỏng vận hành trong dải ~2960–10560 W (phân vị 1–99); dải cũ [3500; 9000] cắt vào vùng an toàn nên độ chính xác chỉ 0,17. Dải mới → độ chính xác 0,79.'],
    ['OSF', 'mòn × mômen > {12800; 13900; 14500}', 'Xác suất hỏng tăng vọt đúng tại các mốc này theo từng loại (L: 0,88; M: 0,94; H: 0,93); ngưỡng tăng dần theo độ cứng sản phẩm.'],
    ['TWF', 'độ mòn dao > 244 phút', 'Xác suất hỏng nhảy bậc từ 0,57 (mốc 240) lên 0,80 (mốc 244); dao không hỏng mòn tối đa chỉ tới 253.'],
], widths=[2.0, 4.4, 10.3])
note('Bằng chứng đóng góp (cô lập).', 'Khi bổ sung bốn đặc trưng này vào mô hình hồi quy logistic cơ sở mà giữ nguyên '
     'mọi thành phần khác, F1 tăng từ 0,231 lên 0,716 và AUC-PR tăng từ 0,220 lên 0,645 (khoảng 2,9 lần). Với bộ '
     'ngưỡng khôi phục, độ chính xác gộp của luật tăng từ 0,26 (cấu hình AI4I gốc) lên 0,81 trên A và giữ 0,82 trên B '
     '— xác nhận biên là bất biến, không phải tinh chỉnh trên tập kiểm thử.')

heading('5.4 Cơ chế ngưỡng và tính bất biến hạng', 3)
para('Hàm bản lề max(0, ·) không đơn thuần mang tính hình thức mà là vị trí duy nhất tại đó giá trị ngưỡng thực sự '
     'tác động đến mô hình cây. Hình 2 minh họa hai cơ chế liên quan.')
figure(ASSETS+'07_hinge.png', 'Hình 2. ', '(a) Hàm bản lề nén phẳng vùng an toàn về 0; thay đổi ngưỡng làm thay đổi tập mẫu bị nén. (b) Phép trừ hằng số bảo toàn thứ tự nên mô hình cây cho kết quả không đổi.', 6.3)
table('Bảng 7. ', 'Tác động của thay đổi ngưỡng theo loại đặc trưng.', ['Loại đặc trưng', 'Ngưỡng có tác động?', 'Cơ chế'], [
    ['Có hàm bản lề (nguy_tan_nhiet, lech_cong_suat, mon_twf)', 'Có', 'Thay đổi ngưỡng làm thay đổi tập mẫu bị nén về 0, dẫn tới thay đổi thứ hạng'],
    ['Không hàm bản lề (bien_overstrain)', 'Không', 'Phép trừ hằng số bảo toàn thứ tự; mô hình cây chỉ phụ thuộc thứ hạng nên kết quả không đổi (kiểm chứng: giảm ngưỡng 100–200 đơn vị, F1 giữ nguyên 0,783)'],
], widths=[6.0, 3.0, 7.7])

heading('5.5 Lựa chọn họ mô hình', 3)
para('Do tín hiệu phi tuyến và tập trung ở các đuôi phân phối, mô hình tuyến tính thuần tỏ ra hạn chế; thực nghiệm '
     'xác nhận mô hình cây vượt trội (F1 tăng từ 0,716 lên 0,783). Nghiên cứu sử dụng ba mô hình cây theo ba cơ chế '
     'khác nhau: Random Forest [5] (đóng bao, giảm phương sai), XGBoost [6] (tăng cường, giảm độ chệch) và ExtraTrees '
     '[7] (ngẫu nhiên hóa ngưỡng). Mô hình hồi quy logistic được giữ lại trong tổ hợp mặc dù hiệu năng đơn lẻ thấp '
     'hơn, do mô hình cây không ngoại suy được ngoài miền giá trị của A; hồi quy logistic ngoại suy có hướng nên bù '
     'đắp cho khoảng 2,8% số mẫu B vượt ngoài miền quan sát của A. Học sâu không được sử dụng vì với chỉ 7 biến và '
     '14.000 mẫu, mạng nơ-ron có dung lượng dư thừa, dễ quá khớp phân phối A và khó diễn giải.')

heading('5.6 Hiệu chỉnh dịch chuyển bằng trọng số tỉ lệ mật độ', 3)
para('Sau khi thiết kế đặc trưng đã hấp thụ phần lớn dịch chuyển, một lớp hiệu chỉnh thống kê được bổ sung: mỗi mẫu '
     'huấn luyện được gán trọng số theo mức độ tương đồng với phân phối B. Tỉ lệ mật độ được ước lượng thông qua một '
     'bộ phân loại dịch chuyển (gán nhãn giả A = 0, B = 1): w(x) = P(B|x) / P(A|x) = p / (1 − p).')
figure(ASSETS+'08_reweight.png', 'Hình 3. ', 'Hiệu chỉnh trọng số điều chỉnh phân phối huấn luyện của A về phía B mà không sử dụng nhãn B.', 5.6)
para('Trọng số được cắt ngưỡng trong khoảng [0,2; 10] và chuẩn hóa về trung bình bằng 1, vì một số ít mẫu hiếm của A '
     'nằm trong vùng B có xác suất xấp xỉ 1, khiến trọng số thô tăng đến khoảng 77 và một mẫu đơn lẻ chi phối toàn bộ '
     'gradient; việc cắt ngưỡng giữ ổn định phương sai (chỉ 0,6% số mẫu chạm giới hạn).')
note('Về mức cải thiện nhỏ của hiệu chỉnh.', 'F1 giữ nguyên ở 0,783 và AUC-PR gần như không đổi (0,675 → 0,671) là '
     'kết quả nhất quán với lý thuyết. Drift-AUC tính riêng trên các đặc trưng biên chỉ đạt 0,53 (Hình 7), nghĩa là '
     'trên không gian đặc trưng này A và B gần như không phân biệt được, do đó không còn nhiều dịch chuyển để hiệu '
     'chỉnh. Điều này khẳng định rằng thiết kế đặc trưng vật lý đúng đắn làm giảm nhu cầu hiệu chỉnh bằng thống kê.')

heading('5.7 Kiểm định có trọng số quan trọng (IWV)', 3)
para('Ràng buộc trọng tâm là mô hình được chấm trên B nhưng không được sử dụng nhãn B để lựa chọn. Cơ sở lý thuyết '
     'của IWV [3] nằm ở việc kỳ vọng hàm mất mát trên B có thể biểu diễn thành kỳ vọng có trọng số trên A: '
     '𝔼_B[L] = 𝔼_A[ w(x) · L ], với w(x) = p_B(x) / p_A(x). Quy trình gồm: (i) xác suất ngoài nếp gấp (out-of-fold) '
     'trên A; (ii) các độ đo Precision, Recall, F1 có trọng số; (iii) quét ngưỡng để cực đại hóa F1 có trọng số. '
     'Toàn bộ quá trình lựa chọn mô hình, ngưỡng và trọng số bỏ phiếu đều dựa trên IWV; nhãn B chỉ được chấm một lần '
     'duy nhất ở bước cuối. Cỡ mẫu hiệu dụng chỉ đạt 24,2% khiến ước lượng IWV có nhiễu; vì vậy không nên tin cậy các '
     'chênh lệch IWV nhỏ hơn khoảng 0,005.')

heading('5.8 Học kết hợp', 3)
table('Bảng 8. ', 'So sánh hai phương án học kết hợp.', ['Phương án', 'Kết quả', 'Đánh giá'], [
    ['Bỏ phiếu — trung bình có trọng số (IWV chọn RF 0,25 / ExtraTrees 0,25 / XGB 0,5)', 'F1-IWV = 0,800; F1-B = 0,783', 'Được chọn — đơn giản, ít tham số, ít quá khớp dịch chuyển'],
    ['Xếp chồng — mô hình tổng hợp logistic (thiên về XGB, hệ số 2,6)', 'F1-IWV = 0,800; F1-B = 0,783', 'Bị loại — dồn trọng số vào một mô hình cơ sở, kém ổn định hơn'],
], widths=[7.5, 4.2, 5.0])
para('Hai phương án cho F1 trên B bằng nhau (0,783) và F1-IWV cũng bằng nhau (0,800), tức nằm trong sai số do cỡ mẫu '
     'hiệu dụng thấp, nên được xem là gần tương đương. Trong trường hợp đó, phương án đơn giản và ổn định hơn được ưu '
     'tiên theo nguyên lý tiết kiệm.')

heading('5.9 Ngưỡng quyết định', 3)
para('Ngưỡng mặc định 0,5 chỉ tối ưu cho lớp cân bằng và phân phối cố định — cả hai điều kiện đều không thỏa mãn ở '
     'đây. Phương pháp quét ngưỡng để cực đại hóa F1 trên thang IWV (mô phỏng B), thu được giá trị 0,585. Ngưỡng này '
     'gần với ngưỡng tối ưu lý tưởng ước lượng trực tiếp trên B, cho thấy IWV đáng tin cậy; đồng thời đỉnh của đường '
     'F1 theo ngưỡng khá phẳng quanh 0,585, cho thấy lựa chọn ngưỡng có tính ổn định.')

# ================= 6. KET QUA =================
pagebreak()
heading('6. Thực nghiệm và kết quả', 2)
figure(ASSETS+'01_journey.png', 'Hình 4. ', 'Diễn tiến cải thiện qua sáu phiên bản. Mức cải thiện lớn nhất (v0 → v1) đến từ thiết kế đặc trưng vật lý; từ v2 mọi phiên bản hội tụ về 0,783 (chạm trần dữ liệu).', 5.6)
table('Bảng 9. ', 'Kết quả mô hình cuối cùng v6 (bỏ phiếu, ngưỡng 0,585) trên Dây chuyền B.', ['Độ đo', 'Giá trị'], [
    ['F1', '0,783'], ['AUC-ROC', '0,872'], ['AUC-PR', '0,666'], ['Precision', '0,817'], ['Recall', '0,751'],
], widths=[4.0, 4.0])
para('Trong 477 thiết bị thực tế hỏng trên B, mô hình phát hiện 358 (dương thật) và bỏ sót 119 (âm giả); trong 438 '
     'cảnh báo phát ra, 80 là báo động giả (dương giả). Tỉ lệ cảnh báo đúng xấp xỉ bốn trên năm.')
figure(ASSETS+'05_confusion.png', 'Hình 5. ', 'Ma trận nhầm lẫn của mô hình v6 trên Dây chuyền B.', 3.6)
figure(ASSETS+'03_psi.png', 'Hình 6. ', 'Chỉ số PSI: dịch chuyển tập trung ở biến thô; bốn đặc trưng biên có PSI xấp xỉ 0.', 5.6)
figure(ASSETS+'06_drift.png', 'Hình 7. ', 'Drift-AUC: đặc trưng biên không phân biệt được A và B (0,53 ≈ dự đoán ngẫu nhiên).', 5.6)
figure(ASSETS+'04_invariance.png', 'Hình 8. ', 'Xác suất hỏng có điều kiện P(hỏng | vượt ngưỡng) gần như không đổi khi chuyển từ A sang B, xác nhận cơ chế sinh lỗi bất biến.', 4.6)
heading('6.1 Đánh giá độ tin cậy bằng bootstrap', 3)
para('Ước lượng điểm F1 = 0,783 với khoảng tin cậy 95% là [0,753; 0,811], ước lượng bằng phương pháp bootstrap [8]. '
     'Khoảng rộng (± 0,03) phản ánh việc chỉ có 477 mẫu hỏng trên B; các khoảng tin cậy của những phiên bản đầu bảng '
     'chồng lấn nhau, do đó không thể kết luận hơn kém dựa trên ước lượng điểm đơn lẻ. So sánh theo cặp trên cùng mẫu '
     'bootstrap cho kết quả tin cậy hơn: xác suất F1 của v6 lớn hơn F1 của XGBoost đạt 100%, với chênh lệch trung '
     'bình +0,015 và khoảng tin cậy 95% [+0,009; +0,022] nằm hoàn toàn ở phía dương — học kết hợp vượt trội nhất quán '
     'về hướng, tuy biên độ nhỏ.')

# ================= 7. THAO LUAN =================
heading('7. Thảo luận: hàm ý vận hành và bảo trì', 2)
runs_para([('Độ mòn dao là chỉ báo then chốt cho bảo trì. ', dict(name=BODY, size=11, bold=True, color=NAVY)),
    ('Đây vừa là tín hiệu dự báo mạnh nhất, vừa ổn định qua dịch chuyển (PSI ≈ 0), lại là biến can thiệp được. Ưu '
     'tiên hàng đầu cho đội bảo trì là giám sát và thay dao theo đúng chu kỳ; luật TWF khôi phục được (mòn > 244) cho '
     'thẳng ngưỡng thay dao.', dict(name=BODY, size=11))])
runs_para([('Cảnh báo theo ngưỡng vật lý áp dụng xuyên nhà máy. ', dict(name=BODY, size=11, bold=True, color=NAVY)),
    ('Vì xác suất hỏng có điều kiện gần như không đổi khi chuyển A → B, có thể thiết lập cảnh báo theo ngưỡng ngay cả '
     'khi thay đổi thiết bị hoặc dây chuyền mà không cần huấn luyện lại.', dict(name=BODY, size=11))])
runs_para([('Điều chỉnh ngưỡng theo chi phí thực tế. ', dict(name=BODY, size=11, bold=True, color=NAVY)),
    ('Do bỏ sót đắt hơn báo động giả, có thể hạ ngưỡng để tăng độ nhạy khi chi phí dừng máy thấp.', dict(name=BODY, size=11))])
table('Bảng 10. ', 'Khuyến nghị hai cấp cảnh báo.', ['Cấp', 'Ngưỡng', 'Ưu tiên', 'Hành động vận hành'], [
    ['Cấp 1 (cao)', 'Ngưỡng cao', 'Precision', 'Dừng máy kiểm tra ngay; độ tin cậy cao, ít báo nhầm'],
    ['Cấp 2 (thấp)', 'Ngưỡng thấp', 'Recall', 'Tăng cường theo dõi, lập lịch bảo trì gần; phát hiện thêm thiết bị nghi ngờ'],
], widths=[2.2, 2.4, 2.2, 9.9])

# ================= 8. HAN CHE =================
heading('8. Hạn chế của phương pháp', 2)
table('Bảng 11. ', 'Các hạn chế và ảnh hưởng.', ['Hạn chế', 'Bản chất và ảnh hưởng'], [
    ['Mô hình cây không ngoại suy', 'Khoảng 2,8% số mẫu B nằm ngoài miền giá trị của A, dẫn tới dự đoán phẳng; hồi quy logistic chỉ bù đắp một phần.'],
    ['Giả định dịch chuyển hiệp biến chưa được chứng minh trực tiếp', 'Hiệu chỉnh trọng số và IWV chỉ đúng nếu P(hỏng | đặc trưng) bất biến; không có nhãn B để kiểm chứng. Nếu tồn tại dịch chuyển khái niệm thực sự, tỉ lệ mật độ mất hiệu lực.'],
    ['Độ tin cậy của IWV bị giới hạn', 'Cỡ mẫu hiệu dụng 24,2% khiến ước lượng nhiễu; chênh lệch giữa bỏ phiếu và xếp chồng nằm trong sai số.'],
    ['Không gian cải thiện còn hẹp', 'Thiết kế đặc trưng đã hấp thụ phần lớn dịch chuyển, nên dư địa cải thiện thêm bằng kỹ thuật xử lý dịch chuyển là hạn chế.'],
    ['Trần hiệu năng của dữ liệu', 'Sau khi khôi phục luật, khoảng 25% ca hỏng vẫn không khớp bất kỳ luật nào — tương ứng lỗi ngẫu nhiên (Random Failures của AI4I [1]). Đây là nhiễu bất khả quy, giới hạn trần F1 ở mức xấp xỉ 0,78, khớp với việc mọi phiên bản từ v2 hội tụ về 0,783.'],
], widths=[5.0, 11.7])

# ================= 9. HUONG PHAT TRIEN =================
heading('9. Hướng phát triển', 2)
note('Cải tiến đã thực hiện.', 'Hướng "hiệu chỉnh lại hằng số ngưỡng" nêu ở phiên bản trước đã được triển khai: bộ '
     'ngưỡng PWF/OSF được khôi phục từ Dây chuyền A và bổ sung luật TWF (Mục 5.3), nâng độ chính xác gộp của luật từ '
     '0,26 lên 0,81 và đưa F1 từ 0,781 lên 0,783. Các hướng còn lại:')
for i, t in enumerate([
    'Xây dựng đặc trưng theo thời gian nhằm nắm bắt xu hướng mòn dao thay vì trạng thái tức thời, cho phép cảnh báo sớm hơn.',
    'Áp dụng các kỹ thuật thích ứng miền nâng cao (CORAL, căn chỉnh đặc trưng đối kháng) để xử lý phần dịch chuyển còn sót ở đuôi phân phối.',
    'Thu thập một lượng nhỏ nhãn trên B thông qua học chủ động nhằm kiểm chứng trực tiếp giả định dịch chuyển hiệp biến.',
    'Áp dụng dự đoán bảo giác (conformal prediction) để cung cấp khoảng tin cậy cho mỗi cảnh báo.',
], 1):
    p = para(justify=True, space_after=3); p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(f'{i}.  '); _set_font(r, BODY, 11, True)
    r2 = p.add_run(t); _set_font(r2, BODY, 11)

# ================= 10. KET LUAN =================
heading('10. Kết luận', 2)
para('Nghiên cứu giải quyết hai thành phần trọng số cao của tiêu chí đánh giá — xử lý dịch chuyển và kết quả trên tập '
     'kiểm thử — thông qua một luận điểm xuyên suốt: xây dựng đặc trưng trên các hằng số cơ chế bất biến (khôi phục từ '
     'A, kiểm chứng trên B) thay vì hiệu chỉnh một mô hình đã học ranh giới bị dịch chuyển. Nhờ đó, ranh giới quyết '
     'định học trên Dây chuyền A chuyển giao trực tiếp sang Dây chuyền B; các bước hiệu chỉnh trọng số, hiệu chỉnh '
     'ngưỡng và học kết hợp chỉ đóng vai trò tinh chỉnh cuối cùng, toàn bộ được lựa chọn bằng kiểm định có trọng số '
     'quan trọng mà không rò rỉ nhãn tập kiểm thử. Kết quả cuối cùng đạt F1 = 0,783 trên Dây chuyền B, cải thiện 3,4 '
     'lần so với mô hình cơ sở, và phân tích cho thấy con số này đã chạm trần hiệu năng mà dữ liệu cho phép.')

# ================= PHU LUC =================
heading('Phụ lục. Kiểm soát rò rỉ dữ liệu', 2)
table('Bảng 12. ', 'Nguy cơ rò rỉ và biện pháp phòng ngừa.', ['Nguy cơ', 'Biểu hiện', 'Biện pháp phòng ngừa'], [
    ['Rò rỉ dữ liệu', 'Điểm số trên tập kiểm thử cao bất thường', 'Chuẩn hóa, mã hóa, chọn đặc trưng chỉ huấn luyện trên tập huấn luyện, đóng gói trong Pipeline; quan hệ cơ sở–tổng hợp dùng dự đoán ngoài nếp gấp'],
    ['Rò rỉ nhãn kiểm thử', 'Lựa chọn cấu hình theo điểm số trên B', 'Lựa chọn mô hình, ngưỡng và trọng số chỉ bằng IWV; nhãn B chấm một lần cuối'],
    ['Mất cân bằng lớp', 'Độ chính xác 92%, độ nhạy bằng 0', 'Loại độ chính xác; dùng F1/AUC-PR, trọng số lớp và hiệu chỉnh ngưỡng'],
    ['Quá khớp', 'Sai số huấn luyện gần 0 nhưng kém trên kiểm thử', 'Kiểm định chéo chọn tham số; chính quy hóa; học kết hợp; đặc trưng bất biến'],
], widths=[3.2, 4.5, 9.0])

# ================= TAI LIEU THAM KHAO =================
heading('Tài liệu tham khảo', 2)
REFS = [
    '[1] S. Matzka, "Explainable Artificial Intelligence for Predictive Maintenance Applications," 2020 Third International Conference on Artificial Intelligence for Industries (AI4I), 2020, pp. 69–74. (Bộ dữ liệu AI4I 2020 Predictive Maintenance, UCI Machine Learning Repository.)',
    '[2] H. Shimodaira, "Improving predictive inference under covariate shift by weighting the log-likelihood function," Journal of Statistical Planning and Inference, vol. 90, no. 2, pp. 227–244, 2000.',
    '[3] M. Sugiyama, M. Krauledat, and K.-R. Müller, "Covariate shift adaptation by importance weighted cross validation," Journal of Machine Learning Research, vol. 8, pp. 985–1005, 2007.',
    '[4] J. Quiñonero-Candela, M. Sugiyama, A. Schwaighofer, and N. D. Lawrence (Eds.), Dataset Shift in Machine Learning. Cambridge, MA: MIT Press, 2009.',
    '[5] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.',
    '[6] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," Proc. 22nd ACM SIGKDD, pp. 785–794, 2016.',
    '[7] P. Geurts, D. Ernst, and L. Wehenkel, "Extremely Randomized Trees," Machine Learning, vol. 63, no. 1, pp. 3–42, 2006.',
    '[8] B. Efron and R. J. Tibshirani, An Introduction to the Bootstrap. New York: Chapman & Hall/CRC, 1993.',
    '[9] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, "SMOTE: Synthetic Minority Over-sampling Technique," Journal of Artificial Intelligence Research, vol. 16, pp. 321–357, 2002.',
    '[10] T. Saito and M. Rehmsmeier, "The Precision-Recall Plot Is More Informative than the ROC Plot When Evaluating Binary Classifiers on Imbalanced Datasets," PLoS ONE, vol. 10, no. 3, e0118432, 2015.',
    '[11] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.',
]
for ref in REFS:
    p = para(justify=True, space_after=3, size=10)
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)
    r = p.add_run(ref); _set_font(r, SANS, 10)

OUT = 'bao_cao/BAO_CAO_PredMaintenance.docx'
doc.save(OUT)
print('DA LUU:', OUT, '| so doan:', len(doc.paragraphs), '| so bang:', len(doc.tables))
